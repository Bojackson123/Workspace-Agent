"""RFI file tools for the Action MCP — parse and fill customer templates.

These tools let the RFI Response Engine work with a customer's *native*
``.xlsx`` / ``.docx`` RFI document instead of converting it to a Google-native
format, so the filled result is byte-for-byte the customer's own template.

Three tools, all keyed by a Shared-Drive file id:

* ``upload_binary_file`` — store base64 bytes (a Chat attachment the backend
  downloaded) as a new file on the Shared Drive; returns its id.
* ``extract_rfi_questions`` — download the file and pull out the list of
  questions, recording for each the exact location its answer must be written
  back to (a spreadsheet cell or a Word table cell / paragraph anchor).
* ``fill_rfi_answers`` — download the file, write each answer into its recorded
  location, and upload the result as a new "… — Sanmina Response" file.

Unlike the Docs/Sheets tools, these return **JSON strings** because they are
driven deterministically by the backend (parsed in Python), not narrated to an
LLM.
"""

from __future__ import annotations

import base64
import io
import json
import logging
import os
import re
from typing import Final

from googleapiclient.errors import HttpError
from googleapiclient.http import MediaIoBaseDownload, MediaIoBaseUpload
from mcp.server.fastmcp import FastMCP
from openpyxl import load_workbook
from openpyxl.utils.cell import coordinate_to_tuple
from docx import Document
from docx.oxml.ns import qn
from docx.table import _Cell
from docx.text.paragraph import Paragraph

from clients import drive as drive_service
from config import settings

from ._retry import retry_on_transient
from .drive import _resolve_parent  # reuse Shared-Drive parent resolution

log = logging.getLogger(__name__)

XLSX_MIME: Final = (
    "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
)
DOCX_MIME: Final = (
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
)

# Caps so the structure dump stays within the parser LLM's context budget.
_MAX_CELLS_PER_SHEET: Final = 600
_MAX_PARAGRAPHS: Final = 400
_MAX_TABLE_ROWS: Final = 200
_MAX_VALUE_LEN: Final = 500


# ── Drive binary helpers ───────────────────────────────────────────────────


def _download_bytes(file_id: str) -> bytes:
    """Download a Shared-Drive file's raw bytes."""
    request = drive_service().files().get_media(
        fileId=file_id, supportsAllDrives=True
    )
    buffer = io.BytesIO()
    downloader = MediaIoBaseDownload(buffer, request)
    done = False
    while not done:
        _status, done = retry_on_transient(downloader.next_chunk)
    return buffer.getvalue()


def _file_meta(file_id: str) -> dict:
    """Return ``{name, mimeType}`` for a Shared-Drive file."""
    return drive_service().files().get(
        fileId=file_id, fields="name, mimeType", supportsAllDrives=True
    ).execute()


def _upload_bytes(
    name: str, mime_type: str, data: bytes, parent_folder_id: str | None
) -> dict:
    """Create a new Shared-Drive file from *data*; return id/name/link dict."""
    media = MediaIoBaseUpload(io.BytesIO(data), mimetype=mime_type, resumable=False)
    created = retry_on_transient(lambda: drive_service().files().create(
        body={"name": name, "parents": [_resolve_parent(parent_folder_id)]},
        media_body=media,
        fields="id, name, webViewLink",
        supportsAllDrives=True,
    ).execute())
    return {
        "file_id": created.get("id"),
        "name": created.get("name"),
        "link": created.get("webViewLink"),
    }


def _find_file_in_parent(name: str, parent_folder_id: str | None) -> str | None:
    """Return the id of an existing non-trashed file named *name* under the parent.

    Lets ``fill_rfi_answers`` be idempotent: repeated calls (a client transport
    retry, or a pipeline re-run) converge onto the one "… — Sanmina Response"
    file instead of creating a fresh duplicate each time.
    """
    parent = _resolve_parent(parent_folder_id)
    escaped = name.replace("\\", "\\\\").replace("'", "\\'")
    query = f"name = '{escaped}' and '{parent}' in parents and trashed = false"
    resp = retry_on_transient(lambda: drive_service().files().list(
        q=query,
        fields="files(id)",
        corpora="allDrives",
        includeItemsFromAllDrives=True,
        supportsAllDrives=True,
        pageSize=1,
    ).execute())
    files = resp.get("files") or []
    return files[0]["id"] if files else None


def _update_bytes(file_id: str, mime_type: str, data: bytes) -> dict:
    """Overwrite an existing Shared-Drive file's content; return id/name/link."""
    media = MediaIoBaseUpload(io.BytesIO(data), mimetype=mime_type, resumable=False)
    updated = retry_on_transient(lambda: drive_service().files().update(
        fileId=file_id,
        media_body=media,
        fields="id, name, webViewLink",
        supportsAllDrives=True,
    ).execute())
    return {
        "file_id": updated.get("id"),
        "name": updated.get("name"),
        "link": updated.get("webViewLink"),
    }


def _upsert_bytes(
    name: str,
    mime_type: str,
    data: bytes,
    response_file_id: str | None,
    parent_folder_id: str | None,
) -> dict:
    """Create-or-overwrite the response file so repeated fills don't duplicate.

    Prefers the caller-supplied *response_file_id* (the file a prior fill
    produced); falls back to a name lookup in the parent, then to a fresh
    upload. A stale *response_file_id* (file deleted → 404) degrades to the
    lookup/create path rather than failing.
    """
    if response_file_id:
        try:
            return _update_bytes(response_file_id, mime_type, data)
        except HttpError as exc:
            if exc.resp.status != 404:
                raise
            log.warning(
                "rfi.fill: response file %s missing; recreating", response_file_id
            )
    existing_id = _find_file_in_parent(name, parent_folder_id)
    if existing_id:
        return _update_bytes(existing_id, mime_type, data)
    return _upload_bytes(name, mime_type, data, parent_folder_id)


def _is_xlsx(name: str, mime: str) -> bool:
    return name.lower().endswith(".xlsx") or "spreadsheet" in (mime or "")


def _is_docx(name: str, mime: str) -> bool:
    return name.lower().endswith(".docx") or "wordprocessing" in (mime or "")


# ── upload_binary_file ─────────────────────────────────────────────────────


def upload_binary_file(
    name: str,
    mime_type: str,
    content_b64: str,
    parent_folder_id: str | None = None,
) -> str:
    """Upload base64-encoded *content_b64* as a new Shared-Drive file.

    Used by the backend to persist a Chat attachment (the customer's RFI) so the
    other RFI tools can read and rewrite it. Returns a JSON object with the new
    ``file_id``, ``name``, and ``link``.
    """
    try:
        data = base64.b64decode(content_b64)
    except Exception as exc:  # noqa: BLE001
        return json.dumps({"error": f"Invalid base64 content: {exc}"})
    try:
        result = _upload_bytes(name, mime_type, data, parent_folder_id)
    except (HttpError, OSError) as exc:
        return json.dumps({"error": f"Upload failed: {exc}"})
    return json.dumps(result)


# ── dump_rfi_structure ─────────────────────────────────────────────────────


def _dump_xlsx(data: bytes) -> dict:
    """Dump every non-empty cell of every sheet with its real A1 address."""
    wb = load_workbook(io.BytesIO(data), data_only=True)
    sheets: list[dict] = []
    for ws in wb.worksheets:
        cells: list[dict] = []
        truncated = False
        for row in ws.iter_rows():
            for cell in row:
                if cell.value is None:
                    continue
                val = str(cell.value).strip()
                if not val:
                    continue
                cells.append({
                    "addr": f"{ws.title}!{cell.coordinate}",
                    "value": val[:_MAX_VALUE_LEN],
                })
            if len(cells) >= _MAX_CELLS_PER_SHEET:
                truncated = True
                break
        sheets.append({
            "sheet": ws.title,
            "max_row": ws.max_row,
            "max_col": ws.max_column,
            "cells": cells[:_MAX_CELLS_PER_SHEET],
            "truncated": truncated,
        })
    wb.close()
    return {"kind": "xlsx", "sheets": sheets}


def _dump_docx(data: bytes) -> dict:
    """Dump Word tables (with row/col indices) and non-empty paragraphs."""
    doc = Document(io.BytesIO(data))
    tables: list[dict] = []
    for t_idx, table in enumerate(doc.tables):
        rows: list[list[str]] = []
        for r_idx, row in enumerate(table.rows):
            if r_idx >= _MAX_TABLE_ROWS:
                break
            rows.append([c.text.strip()[:_MAX_VALUE_LEN] for c in row.cells])
        tables.append({
            "table_index": t_idx,
            "n_rows": len(table.rows),
            "n_cols": len(rows[0]) if rows else 0,
            "rows": rows,
        })
    paragraphs: list[dict] = []
    for p_idx, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if not text:
            continue
        paragraphs.append({"index": p_idx, "text": text[:_MAX_VALUE_LEN]})
        if len(paragraphs) >= _MAX_PARAGRAPHS:
            break
    return {"kind": "docx", "tables": tables, "paragraphs": paragraphs}


def dump_rfi_structure(file_id: str) -> str:
    """Dump an RFI file's raw structure with precise write-back anchors.

    Returns JSON describing the spreadsheet cells — each with its real address
    (e.g. ``"Responses!B5"``) — or the Word tables/paragraphs (with table, row,
    and column indices). An LLM parser reads this to identify the questions and
    choose, FROM THESE REAL ANCHORS, where each answer should be written, so
    write-back stays reliable across arbitrarily-structured customer RFIs.

    Anchor formats ``fill_rfi_answers`` expects back:
      - spreadsheet cell: ``"<SheetName>!<A1>"``   (e.g. ``"Responses!C7"``)
      - Word table cell:  ``"tbl-<t>!r<row>c<col>"`` (0-based indices)
      - Word paragraph:   ``"para-<index>"``        (answer inserted next line)
    """
    try:
        meta = _file_meta(file_id)
        data = _download_bytes(file_id)
    except HttpError as exc:
        return json.dumps({"error": f"Could not read file: {exc}"})

    name, mime = meta.get("name", ""), meta.get("mimeType", "")
    try:
        if _is_xlsx(name, mime):
            structure = _dump_xlsx(data)
        elif _is_docx(name, mime):
            structure = _dump_docx(data)
        else:
            return json.dumps({
                "error": f"Unsupported file type: {name!r} ({mime}). "
                         "Attach an .xlsx or .docx RFI."
            })
    except Exception as exc:  # noqa: BLE001 — surface parse failures as JSON
        log.exception("rfi.dump failed for %s", file_id)
        return json.dumps({"error": f"Failed to read {name!r}: {exc}"})

    structure["file_name"] = name
    return json.dumps(structure)


# ── fill_rfi_answers ───────────────────────────────────────────────────────


# Spreadsheet/Word cells store literal text — they cannot render markdown — so
# any markdown the research LLM emits would show verbatim (e.g. "**ISO 9001**",
# "- item"). Strip the common syntax to clean prose before writing it back.
_MD_CODE_FENCE = re.compile(r"^\s*```.*$", re.MULTILINE)
_MD_LINK = re.compile(r"\[([^\]]+)\]\(([^)\s]+)(?:\s+\"[^\"]*\")?\)")
_MD_IMAGE = re.compile(r"!\[([^\]]*)\]\([^)]*\)")
_MD_EMPHASIS = re.compile(r"(\*\*\*|\*\*|\*|___|__|_|`)(?=\S)(.+?)(?<=\S)\1")
_MD_HEADING = re.compile(r"^\s{0,3}#{1,6}\s+", re.MULTILINE)
_MD_BLOCKQUOTE = re.compile(r"^\s{0,3}>\s?", re.MULTILINE)
_MD_BULLET = re.compile(r"^(\s*)[-*+]\s+", re.MULTILINE)
_MD_HRULE = re.compile(r"^\s{0,3}([-*_])(?:\s*\1){2,}\s*$", re.MULTILINE)


def _md_to_plain(text: str) -> str:
    """Convert markdown the LLM may emit into plain text for a cell.

    Cheap, dependency-free, and conservative: it removes emphasis/heading/quote
    markers, unwraps links to ``text (url)``, and normalises bullets to "• " so
    list answers stay readable. Leaves ordinary prose untouched.
    """
    if not text or not isinstance(text, str):
        return text or ""
    out = _MD_HRULE.sub("", text)
    out = _MD_CODE_FENCE.sub("", out)
    out = _MD_IMAGE.sub(r"\1", out)
    # "[Sanmina](https://…)" → "Sanmina (https://…)"; bare "[label](label)"
    # (text == url) collapses to just the label.
    out = _MD_LINK.sub(
        lambda m: m.group(1) if m.group(1) == m.group(2) else f"{m.group(1)} ({m.group(2)})",
        out,
    )
    out = _MD_HEADING.sub("", out)
    out = _MD_BLOCKQUOTE.sub("", out)
    out = _MD_BULLET.sub(r"\1• ", out)
    # Emphasis last, after structural markers are gone. Run twice to catch
    # nested/adjacent runs (e.g. "**bold _and_ italic**").
    out = _MD_EMPHASIS.sub(r"\2", out)
    out = _MD_EMPHASIS.sub(r"\2", out)
    # Collapse the blank lines left by stripped fences/rules.
    out = re.sub(r"\n{3,}", "\n\n", out)
    return out.strip()


def _fill_xlsx(data: bytes, answers: list[dict]) -> bytes:
    """Write answers into a workbook (formulas/formatting preserved) → bytes."""
    wb = load_workbook(io.BytesIO(data))  # data_only=False keeps formulas
    for ans in answers:
        location = ans.get("location") or ""
        sheet, _, coord = location.rpartition("!")
        if not sheet or sheet not in wb.sheetnames:
            continue
        ws = wb[sheet]
        try:
            row, col = coordinate_to_tuple(coord)
        except Exception:  # noqa: BLE001
            continue
        ws.cell(row=row, column=col, value=_md_to_plain(ans.get("answer", "")))
    out = io.BytesIO()
    wb.save(out)
    return out.getvalue()


def _insert_paragraph_after(paragraph: Paragraph, text: str) -> None:
    """Insert a new paragraph containing *text* immediately after *paragraph*."""
    new_p = paragraph._p.makeelement(qn("w:p"), {})
    paragraph._p.addnext(new_p)
    Paragraph(new_p, paragraph._parent).add_run(text)


def _fill_docx(data: bytes, answers: list[dict]) -> bytes:
    """Write answers into a Word doc (table cells / after question paras) → bytes."""
    doc = Document(io.BytesIO(data))
    for ans in answers:
        location = ans.get("location") or ""
        answer_text = _md_to_plain(ans.get("answer", ""))
        if location.startswith("tbl-"):
            try:
                tbl_part, cell_part = location.split("!", 1)
                t_idx = int(tbl_part[len("tbl-"):])
                r_idx = int(cell_part[1:cell_part.index("c")])
                c_idx = int(cell_part[cell_part.index("c") + 1:])
            except (ValueError, IndexError):
                continue
            if t_idx >= len(doc.tables):
                continue
            rows = doc.tables[t_idx].rows
            if r_idx >= len(rows) or c_idx >= len(rows[r_idx].cells):
                continue
            cell: _Cell = rows[r_idx].cells[c_idx]
            cell.text = answer_text
        elif location.startswith("para-"):
            try:
                p_idx = int(location[len("para-"):])
            except ValueError:
                continue
            if p_idx < len(doc.paragraphs):
                _insert_paragraph_after(doc.paragraphs[p_idx], answer_text)
    out = io.BytesIO()
    doc.save(out)
    return out.getvalue()


def _response_name(original: str) -> str:
    """Derive the filled-file name: ``RFI.xlsx`` → ``RFI — Sanmina Response.xlsx``."""
    stem, ext = os.path.splitext(original)
    return f"{stem} — Sanmina Response{ext}"


def fill_rfi_answers(
    file_id: str,
    answers_json: str,
    response_file_id: str | None = None,
    parent_folder_id: str | None = None,
) -> str:
    """Write answers into the RFI template and upsert a filled copy.

    Args:
        file_id: Shared-Drive id of the original RFI file.
        answers_json: JSON array of ``{"location": <answer_location>,
            "answer": <text>}`` objects. ``location`` values come from
            ``extract_rfi_questions``' ``answer_location`` field.
        response_file_id: id of a response file a previous fill produced. When
            given, its content is overwritten in place instead of creating a new
            file — so retries/re-runs converge onto one file. Idempotent even
            without it: a same-named response file in the parent is reused.

    Returns a JSON object ``{file_id, name, link, written}``.
    """
    try:
        answers = json.loads(answers_json)
        if not isinstance(answers, list):
            raise ValueError("answers_json must be a JSON array")
    except (json.JSONDecodeError, ValueError) as exc:
        return json.dumps({"error": f"Invalid answers_json: {exc}"})

    try:
        meta = _file_meta(file_id)
        data = _download_bytes(file_id)
    except HttpError as exc:
        return json.dumps({"error": f"Could not read file: {exc}"})

    name, mime = meta.get("name", ""), meta.get("mimeType", "")
    try:
        if _is_xlsx(name, mime):
            filled, out_mime = _fill_xlsx(data, answers), XLSX_MIME
        elif _is_docx(name, mime):
            filled, out_mime = _fill_docx(data, answers), DOCX_MIME
        else:
            return json.dumps({"error": f"Unsupported file type: {name!r} ({mime})."})
    except Exception as exc:  # noqa: BLE001
        log.exception("rfi.fill failed for %s", file_id)
        return json.dumps({"error": f"Failed to fill {name!r}: {exc}"})

    try:
        result = _upsert_bytes(
            _response_name(name), out_mime, filled, response_file_id, parent_folder_id
        )
    except HttpError as exc:
        return json.dumps({"error": f"Upload of filled file failed: {exc}"})
    result["written"] = len(answers)
    return json.dumps(result)


_TOOLS = (
    upload_binary_file,
    dump_rfi_structure,
    fill_rfi_answers,
)


def register(mcp: FastMCP) -> None:
    """Register every RFI file tool defined in this module onto *mcp*."""
    for fn in _TOOLS:
        mcp.tool()(fn)
