"""Word (.docx) structure dump and answer write-back."""

from __future__ import annotations

import io

from docx import Document
from docx.oxml.ns import qn
from docx.table import _Cell
from docx.text.paragraph import Paragraph

from ._limits import _MAX_PARAGRAPHS, _MAX_TABLE_ROWS, _MAX_VALUE_LEN
from ._markdown import _md_to_plain


def _dump_docx(data: bytes) -> dict:
    """Dump Word tables (with row/col indices) and non-empty paragraphs."""
    doc = Document(io.BytesIO(data))
    tables: list[dict] = []
    for t_idx, table in enumerate(doc.tables):
        rows: list[list[str]] = []
        for r_idx, row in enumerate(table.rows):
            if r_idx >= _MAX_TABLE_ROWS:
                break
            rows.append([c.text.strip()[:_MAX_VALUE_LEN] for c in row.cells])
        tables.append({
            "table_index": t_idx,
            "n_rows": len(table.rows),
            "n_cols": len(rows[0]) if rows else 0,
            "rows": rows,
        })
    paragraphs: list[dict] = []
    for p_idx, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if not text:
            continue
        paragraphs.append({"index": p_idx, "text": text[:_MAX_VALUE_LEN]})
        if len(paragraphs) >= _MAX_PARAGRAPHS:
            break
    return {"kind": "docx", "tables": tables, "paragraphs": paragraphs}


def _insert_paragraph_after(paragraph: Paragraph, text: str) -> None:
    """Insert a new paragraph containing *text* immediately after *paragraph*."""
    new_p = paragraph._p.makeelement(qn("w:p"), {})
    paragraph._p.addnext(new_p)
    Paragraph(new_p, paragraph._parent).add_run(text)


def _fill_docx(data: bytes, answers: list[dict]) -> bytes:
    """Write answers into a Word doc (table cells / after question paras) → bytes."""
    doc = Document(io.BytesIO(data))
    for ans in answers:
        location = ans.get("location") or ""
        answer_text = _md_to_plain(ans.get("answer", ""))
        if location.startswith("tbl-"):
            try:
                tbl_part, cell_part = location.split("!", 1)
                t_idx = int(tbl_part[len("tbl-"):])
                r_idx = int(cell_part[1:cell_part.index("c")])
                c_idx = int(cell_part[cell_part.index("c") + 1:])
            except (ValueError, IndexError):
                continue
            if t_idx >= len(doc.tables):
                continue
            rows = doc.tables[t_idx].rows
            if r_idx >= len(rows) or c_idx >= len(rows[r_idx].cells):
                continue
            cell: _Cell = rows[r_idx].cells[c_idx]
            cell.text = answer_text
        elif location.startswith("para-"):
            try:
                p_idx = int(location[len("para-"):])
            except ValueError:
                continue
            if p_idx < len(doc.paragraphs):
                _insert_paragraph_after(doc.paragraphs[p_idx], answer_text)
    out = io.BytesIO()
    doc.save(out)
    return out.getvalue()
